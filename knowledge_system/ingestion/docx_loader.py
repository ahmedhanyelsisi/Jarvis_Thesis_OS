"""Load text and document metadata from DOCX files with python-docx."""

from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document


def _isoformat(value: datetime | None) -> str | None:
    return value.isoformat() if value is not None else None


class DocxLoader:
    """Extract paragraphs, tables, and core properties from a DOCX file."""

    @staticmethod
    def load(file_path: str | Path) -> dict[str, Any]:
        path = Path(file_path).expanduser()
        if not path.is_file():
            raise FileNotFoundError(f"DOCX document not found: {path}")
        if path.suffix.lower() != ".docx":
            raise ValueError(f"Expected a .docx document, received: {path.name}")

        document = Document(str(path))
        content_parts = [
            paragraph.text.strip()
            for paragraph in document.paragraphs
            if paragraph.text.strip()
        ]
        for table in document.tables:
            for row in table.rows:
                row_text = "\t".join(cell.text.strip() for cell in row.cells).strip()
                if row_text:
                    content_parts.append(row_text)

        properties = document.core_properties
        metadata: dict[str, Any] = {"document_type": "docx"}
        property_values = {
            "author": properties.author,
            "title": properties.title,
            "subject": properties.subject,
            "created": _isoformat(properties.created),
            "modified": _isoformat(properties.modified),
        }
        metadata.update(
            {
                key: str(value).strip()
                for key, value in property_values.items()
                if value is not None and str(value).strip()
            }
        )

        return {
            "filename": path.name,
            "content": "\n".join(content_parts),
            "metadata": metadata,
        }


def load_docx(file_path: str | Path) -> dict[str, Any]:
    """Functional interface for :class:`DocxLoader`."""

    return DocxLoader.load(file_path)


DOCXLoader = DocxLoader
